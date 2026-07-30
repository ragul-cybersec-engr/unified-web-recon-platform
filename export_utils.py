import re
import textwrap
from datetime import datetime
from io import BytesIO
from typing import Tuple
from xml.dom import minidom
from xml.etree import ElementTree as ET

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def _safe_name(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value or "").strip("._-")
    return cleaned or "organization"


def build_recon_report(orgname: str, storage, engine) -> dict:
    if hasattr(storage, "normalize_orgname"):
        org_key = storage.normalize_orgname(orgname)
    else:
        org_key = _safe_name(orgname)

    all_jobs = storage.load_all()
    matches = [
        r
        for r in all_jobs
        if r.get("orgname") == orgname or r.get("org_key") == org_key
    ]
    latest = matches[-1] if matches else {}

    return {
        "organization": orgname,
        "org_key": org_key,
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "scan_type": latest.get("scan_type", "extended"),
        "domains": latest.get("domains", []),
        "subdomains": engine.read_output_list(orgname, "subfinder.txt"),
        "open_ports": engine.read_output_list(orgname, "naabu.txt"),
        "live_hosts": engine.read_output_list(orgname, "live.txt"),
        "history": storage.load_org(orgname),
    }


def _report_lines(report: dict):
    lines = [
        f"Recon Report - {report['organization']}",
        f"Generated At: {report['generated_at']}",
        f"Scan Type: {report.get('scan_type', 'extended')}",
        "",
        "Domains:",
    ]
    domains = report.get("domains", [])
    if domains:
        lines.extend([f"- {d}" for d in domains])
    else:
        lines.append("- No domains available")

    lines.extend(["", f"Subdomains ({len(report.get('subdomains', []))}):"])
    lines.extend([f"- {s}" for s in report.get("subdomains", [])] or ["- None"])

    lines.extend(["", f"Open Ports ({len(report.get('open_ports', []))}):"])
    lines.extend([f"- {p}" for p in report.get("open_ports", [])] or ["- None"])

    lines.extend(["", f"Live Hosts ({len(report.get('live_hosts', []))}):"])
    lines.extend([f"- {h}" for h in report.get("live_hosts", [])] or ["- None"])

    lines.extend(["", "History Events:"])
    history = report.get("history", [])
    if history:
        for event in history:
            evt = ", ".join([f"{k}={v}" for k, v in event.items()])
            lines.append(f"- {evt}")
    else:
        lines.append("- No events")
    return lines


def _to_txt(report: dict) -> bytes:
    return ("\n".join(_report_lines(report)) + "\n").encode("utf-8")


def _to_xml(report: dict) -> bytes:
    root = ET.Element("recon_report")
    ET.SubElement(root, "organization").text = str(report["organization"])
    ET.SubElement(root, "org_key").text = str(report["org_key"])
    ET.SubElement(root, "generated_at").text = str(report["generated_at"])
    ET.SubElement(root, "scan_type").text = str(report.get("scan_type", "extended"))

    domains = ET.SubElement(root, "domains")
    for item in report.get("domains", []):
        ET.SubElement(domains, "domain").text = str(item)

    subdomains = ET.SubElement(root, "subdomains")
    for item in report.get("subdomains", []):
        ET.SubElement(subdomains, "subdomain").text = str(item)

    open_ports = ET.SubElement(root, "open_ports")
    for item in report.get("open_ports", []):
        ET.SubElement(open_ports, "port").text = str(item)

    live_hosts = ET.SubElement(root, "live_hosts")
    for item in report.get("live_hosts", []):
        ET.SubElement(live_hosts, "host").text = str(item)

    history = ET.SubElement(root, "history")
    for event in report.get("history", []):
        event_node = ET.SubElement(history, "event")
        for key, value in event.items():
            ET.SubElement(event_node, str(key)).text = str(value)

    xml_bytes = ET.tostring(root, encoding="utf-8")
    return minidom.parseString(xml_bytes).toprettyxml(indent="  ", encoding="utf-8")


def _to_docx(report: dict) -> bytes:
    doc = Document()
    doc.add_heading(f"Recon Report - {report['organization']}", level=1)
    doc.add_paragraph(f"Generated At: {report['generated_at']}")
    doc.add_paragraph(f"Scan Type: {report.get('scan_type', 'extended')}")

    doc.add_heading("Domains", level=2)
    for item in report.get("domains", []):
        doc.add_paragraph(str(item), style="List Bullet")
    if not report.get("domains"):
        doc.add_paragraph("No domains available")

    doc.add_heading(f"Subdomains ({len(report.get('subdomains', []))})", level=2)
    for item in report.get("subdomains", []):
        doc.add_paragraph(str(item), style="List Bullet")
    if not report.get("subdomains"):
        doc.add_paragraph("None")

    doc.add_heading(f"Open Ports ({len(report.get('open_ports', []))})", level=2)
    for item in report.get("open_ports", []):
        doc.add_paragraph(str(item), style="List Bullet")
    if not report.get("open_ports"):
        doc.add_paragraph("None")

    doc.add_heading(f"Live Hosts ({len(report.get('live_hosts', []))})", level=2)
    for item in report.get("live_hosts", []):
        doc.add_paragraph(str(item), style="List Bullet")
    if not report.get("live_hosts"):
        doc.add_paragraph("None")

    doc.add_heading("History Events", level=2)
    for event in report.get("history", []):
        doc.add_paragraph(", ".join([f"{k}={v}" for k, v in event.items()]), style="List Bullet")
    if not report.get("history"):
        doc.add_paragraph("No events")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _to_pdf(report: dict) -> bytes:
    lines = _report_lines(report)
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    x = 40
    y = height - 40

    pdf.setFont("Helvetica", 10)
    for line in lines:
        wrapped = textwrap.wrap(line, width=105) or [""]
        for piece in wrapped:
            if y <= 40:
                pdf.showPage()
                pdf.setFont("Helvetica", 10)
                y = height - 40
            pdf.drawString(x, y, piece)
            y -= 14

    pdf.save()
    buffer.seek(0)
    return buffer.read()


def build_export_file(orgname: str, report: dict, export_format: str) -> Tuple[bytes, str, str]:
    fmt = (export_format or "").strip().lower()
    safe_org = _safe_name(orgname)
    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")

    if fmt == "txt":
        return _to_txt(report), f"{safe_org}_{timestamp}.txt", "text/plain"
    if fmt == "xml":
        return _to_xml(report), f"{safe_org}_{timestamp}.xml", "application/xml"
    if fmt == "docx":
        return (
            _to_docx(report),
            f"{safe_org}_{timestamp}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if fmt == "pdf":
        return _to_pdf(report), f"{safe_org}_{timestamp}.pdf", "application/pdf"
    raise ValueError("Unsupported export format")
